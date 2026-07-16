"""Convert docs/SOLUTION_DOCUMENT.md to Word (.docx)."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "docs" / "SOLUTION_DOCUMENT.md"
OUT_PATH = ROOT / "docs" / "FinHealth_Card_Solution_Document.docx"


def set_cell_shading(cell, color: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_formatted_runs(paragraph, text: str, base_bold: bool = False) -> None:
    """Parse inline **bold**, *italic*, `code`, and [links](url)."""
    pattern = re.compile(
        r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\))"
    )
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            run.bold = base_bold
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
        elif token.startswith("["):
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if m:
                run = paragraph.add_run(m.group(1))
                run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
                run.underline = True
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.bold = base_bold


def parse_table_row(line: str) -> list[str]:
    line = line.strip().strip("|")
    return [c.strip() for c in line.split("|")]


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\|?[\s\-:|]+\|?$", line.strip()))


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(ncols):
            cell_text = row[j] if j < len(row) else ""
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            add_formatted_runs(p, cell_text, base_bold=(i == 0))
            if i == 0:
                set_cell_shading(cell, "E8F5E9")
    doc.add_paragraph()


def add_code_block(doc: Document, lines: list[str], lang: str = "") -> None:
    label = f" ({lang})" if lang and lang != "mermaid" else ""
    if lang == "mermaid":
        p = doc.add_paragraph()
        run = p.add_run("Architecture / flow diagram (Mermaid — render in GitHub or Mermaid Live Editor):")
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
    doc.add_paragraph()


def convert(md_path: Path, out_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title page
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("FinHealth Card")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x16, 0x65, 0x34)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Solution Document — Finn. Alternative Score System")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("IDBI · New-To-Credit MSME Credit Assessment · PoC July 2026")
    run.font.size = Pt(11)
    run.italic = True

    doc.add_page_break()

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if stripped == "---":
            doc.add_paragraph()
            i += 1
            continue

        # Code fence
        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, code_lines, lang)
            i += 1
            continue

        # Headings
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            heading_text = stripped[level:].strip()
            if level == 1:
                doc.add_heading(heading_text, level=0)
            else:
                doc.add_heading(heading_text, level=min(level, 3))
            i += 1
            continue

        # Table
        if stripped.startswith("|"):
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                if not is_table_separator(lines[i]):
                    table_rows.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, table_rows)
            continue

        # Blockquote
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            content = stripped.lstrip("> ").strip()
            add_formatted_runs(p, content)
            for run in p.runs:
                run.italic = True
            i += 1
            continue

        # Bullet list
        if re.match(r"^[-*]\s+", stripped):
            while i < len(lines) and re.match(r"^[-*]\s+", lines[i].strip()):
                item = re.sub(r"^[-*]\s+", "", lines[i].strip())
                p = doc.add_paragraph(style="List Bullet")
                add_formatted_runs(p, item)
                i += 1
            continue

        # Numbered list
        if re.match(r"^\d+\.\s+", stripped):
            while i < len(lines) and re.match(r"^\d+\.\s+", lines[i].strip()):
                item = re.sub(r"^\d+\.\s+", "", lines[i].strip())
                p = doc.add_paragraph(style="List Number")
                add_formatted_runs(p, item)
                i += 1
            continue

        # Regular paragraph (collect continuation lines)
        para_lines = [stripped]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if (
                not nxt
                or nxt == "---"
                or nxt.startswith("#")
                or nxt.startswith("|")
                or nxt.startswith("```")
                or nxt.startswith(">")
                or re.match(r"^[-*]\s+", nxt)
                or re.match(r"^\d+\.\s+", nxt)
            ):
                break
            para_lines.append(nxt)
            i += 1
        p = doc.add_paragraph()
        add_formatted_runs(p, " ".join(para_lines))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Created: {out_path}")


if __name__ == "__main__":
    md = Path(sys.argv[1]) if len(sys.argv) > 1 else MD_PATH
    out = Path(sys.argv[2]) if len(sys.argv) > 2 else OUT_PATH
    convert(md, out)
